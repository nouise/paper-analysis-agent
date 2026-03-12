"""
Word 文档解析器

使用 python-docx 解析 .docx 文件
"""

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

from . import DocumentParser, ParseError


class DocxParser(DocumentParser):
    """Word 文档解析器 (.docx)"""

    def supports(self, mime_type: str, file_extension: str) -> bool:
        return (file_extension.lower() in ['.docx'] or
                mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'])

    async def parse(self, file_path: str) -> str:
        """解析 Word 文档"""
        if not HAS_DOCX:
            raise ParseError("Word 解析需要安装 python-docx: poetry install")

        try:
            doc = Document(file_path)
            paragraphs = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格文本（可选）
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n".join(paragraphs)
        except Exception as e:
            raise ParseError(f"Word 解析失败: {e}")
